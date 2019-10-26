import os
import smtplib
from io import BytesIO
from datetime import date

from django.views.generic import ListView
from django.shortcuts import render, redirect
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from email import encoders
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase

from ase_site.auth_core.models import User
from ase_site.data.models import Application
from .forms import MakeRequestForm


class ViewAllRequests(ListView):
    template_name = "ase_site/req/templates/posts.html"
    model = Application

    def get_queryset(self):
        qs = super(ViewAllRequests, self).get_queryset()
        return qs


def approve(request, post_id):
    document = Document()
    document.add_picture('root_files/ase_logo.png', width=Inches(1.25))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Дмитровское ш.,2,Москва', 3,)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading('Запрос!', level=1)
    req = Application.objects.get(id=post_id)
    document.add_paragraph(req.post)
    if req.status == 0:
        req.title = req.title+"(обработан)"
        req.status = 1
        req.save()
    else:
        req.save()
    document.save('root_files/test'+post_id+'.docx')
    #fromaddr = "testase@mail.ru"
    fromaddr = "test_mail_temp2018@mail.ru"
    toaddr = "dr.interned@gmail.com"
    msg = MIMEMultipart()
    msg['From'] = fromaddr
    msg['To'] = toaddr
    msg['Subject'] = "ЗАПРОС"
    body = "Вам пришел запрос"
    msg.attach(MIMEText(body, 'plain'))
    filename = "test.docx"
    attachment = open('root_files/test'+post_id+'.docx', "rb")
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= %s" % filename)
    msg.attach(part)
    server = smtplib.SMTP('smtp.mail.ru', 587)
    server.starttls()
    #server.login(fromaddr, "adminase2018")
    server.login(fromaddr, "temp_pass2018")
    text = msg.as_string()
    server.sendmail(fromaddr, toaddr, text)
    server.quit()
    attachment.close()
    #time.sleep(15)
    os.remove("root_files/test"+post_id+".docx")
    return render(request, 'static/static_files/html/box.html', {'values': ['Запрос отправлен']})


def disapprove(request, post_id):
    req = Application.objects.get(id=post_id)
    req.current_level = req.current_level-1
    req.title = "Запрос №"+str(req.id)+"(отклонен)"
    req.save()
    return render(request, 'static/static_files/html/box.html', {'values': ['Запрос отклонены']})


def fill_word(data):
    document = Document('template.docx')
    paragraphs = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)

    for p in paragraphs:
        if p.text in {str(i) for i in range(len(paragraphs))}:
            p.text = data[int(p.text)]
            print(p.text)

    return document


def create_request(request, application_type):
    if request.method == "POST":
        form = MakeRequestForm(request.POST)
        if form.is_valid():
            application, _ = Application.objects.update_or_create(
                application_type=application_type,
                status=2,
                density=form.cleaned_data['density'],
                volume=form.cleaned_data['volume'],
                delivery_date=form.cleaned_data['delivery_date'],
                delivery_time=form.cleaned_data['delivery_time'],
                car=form.cleaned_data['car'],
                manufacturer_org=form.cleaned_data['manufacturer_org'],
                performing_org=form.cleaned_data['performing_org'],
                application_sender=request.user,
                application_receiver=form.cleaned_data['application_receiver']
            )
            application.save()

            # firm_name = current_user.firm_name
            # current_level = current_user.level
            # form.application_receiver = User.objects.filter(firm_name=firm_name).filter(current_level+1)

            # data = []
            # for i in form.fields:
            #     data.append(request.POST.get(str(i)))
            # doc = fill_word(data)
            # form.date = date.today()
            # data = BytesIO()
            # doc.save(data)
            # response = HttpResponse(data.getvalue(),
            #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            # response['Content-Disposition'] = 'attachment; filename="reports.docx"'
            return redirect('/')
            return response
    else:   
        form = MakeRequestForm()
    return render(request, "ase_site/req/templates/index.html", {"form": form})


def create_beton_request(request):
    return create_request(request, 1)


def create_sand_request(request):
    return create_request(request, 2)


def create_PGS_request(request):
    return create_request(request, 3)
